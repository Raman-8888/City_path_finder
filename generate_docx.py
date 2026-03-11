from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)

# ── Helper: set paragraph shading ────────────────────────────────────────────
def shade_paragraph(para, fill="F0F0F0"):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    pPr.append(shd)

# ── Helper: set cell background ───────────────────────────────────────────────
def set_cell_bg(cell, fill):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill)
    tcPr.append(shd)

# ── Helper: add code block paragraph ─────────────────────────────────────────
def add_code(doc, text):
    for line in text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Inches(0.3)
        shade_paragraph(p, "ECECEC")
        run = p.add_run(line if line else " ")
        run.font.name = 'Courier New'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)
    doc.add_paragraph()

# ── Helper: bullet point ──────────────────────────────────────────────────────
def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after = Pt(3)
    # handle bold **text**
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
    return p

# ── Helper: heading ───────────────────────────────────────────────────────────
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    h.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    h.paragraph_format.space_after  = Pt(6)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    return h

# ── Helper: normal paragraph with inline bold ─────────────────────────────────
def add_para(doc, text, center=False, bold_all=False):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            run.bold = True
            run.font.size = Pt(11)
        else:
            run = p.add_run(part)
            run.font.size = Pt(11)
            if bold_all:
                run.bold = True
    return p

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Data Structures and Algorithm")
run.bold = True
run.font.size = Pt(20)
run.font.color.rgb = RGBColor(0x1A, 0x56, 0x8E)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Inter-City Pathfinder")
run.bold = True
run.font.size = Pt(26)
run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("3D Graph Algorithm Visualizer")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
run.italic = True

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

# Submitted to / by table
tbl = doc.add_table(rows=1, cols=2)
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl.style = 'Table Grid'

c1 = tbl.rows[0].cells[0]
c2 = tbl.rows[0].cells[1]

set_cell_bg(c1, "DEEAF1")
set_cell_bg(c2, "DEEAF1")

p1 = c1.paragraphs[0]
p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p1.add_run("Submitted To\n")
r.bold = True; r.font.size = Pt(11)
r2 = p1.add_run("Mr. Sudeep Chowhan")
r2.font.size = Pt(11)

p2 = c2.paragraphs[0]
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p2.add_run("Submitted By\n")
r.bold = True; r.font.size = Pt(11)
r2 = p2.add_run("Priyanshu Singh (12309720)\nSrijan Ghosh (12314625)")
r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# INDEX
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "Index", 1)

index_data = [
    ("S.No", "Topics"),
    ("1", "Problem Statement"),
    ("2", "Problem-Solving Technique"),
    ("3", "Requirements"),
    ("4", "Code"),
    ("5", "Real Time Example"),
]

tbl = doc.add_table(rows=len(index_data), cols=2)
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.LEFT

for i, (sno, topic) in enumerate(index_data):
    row = tbl.rows[i]
    c1, c2 = row.cells[0], row.cells[1]
    fill = "1A568E" if i == 0 else ("F5F9FC" if i % 2 == 0 else "FFFFFF")
    set_cell_bg(c1, fill)
    set_cell_bg(c2, fill)
    for cell, val in ((c1, sno), (c2, topic)):
        p = cell.paragraphs[0]
        r = p.add_run(val)
        r.bold = (i == 0)
        r.font.size = Pt(11)
        if i == 0:
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 1. PROBLEM STATEMENT
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "1. Problem Statement", 1)

add_para(doc, "The study of graph-based pathfinding algorithms is a cornerstone of computer science, yet the abstract nature of these algorithms makes them difficult to grasp through static textbook examples alone. Traditional teaching methods present several critical gaps that this project aims to resolve:")

bullets = [
    ("**Lack of Visual Context:** Algorithms like Dijkstra's and A* operate on graphs described purely with numbers and notation. Without a spatial, visual representation, students struggle to intuitively understand *why* one algorithm outperforms another."),
    ("**Inability to Observe Step-by-Step Execution:** Textbooks show only the final result of a search. There is no mechanism to observe how the algorithm expands node by node — the most critical insight for understanding algorithmic behaviour."),
    ("**No Interactive Comparison:** Understanding the difference between Dijkstra, A*, BFS, and DFS requires running them side by side under identical conditions. Static notes cannot provide this dynamic, real-time comparison."),
    ("**Disconnect Between Theory and Application:** Graph algorithms form the backbone of GPS navigation, game AI, and network routing. Learners lack a tool that bridges the gap between abstract graph theory and these real-world use cases."),
]
for b in bullets:
    add_bullet(doc, b)

add_para(doc, "This project addresses these gaps by providing a **3D interactive browser-based map** that simulates a network of cities connected by weighted roads, allowing users to visually execute and compare four major pathfinding algorithms in real time.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 2. PROBLEM-SOLVING TECHNIQUE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "2. Problem-Solving Technique", 1)

add_para(doc, "To solve the challenge of making graph algorithms visual and interactive, this system employs a **3D Graph Simulation with Step-by-Step Animation Replay** approach. The solution architecture is built upon the following pillars:")

add_heading(doc, "The Graph Lifecycle", 2)

lifecycle = [
    ("**Build:**", "Constructing the graph with 80 nodes (cities) and weighted edges (roads) using an Adjacency List."),
    ("**Search:**", "Running the selected algorithm to explore paths from start to destination."),
    ("**Animate:**", "Replaying the algorithm's traversal visually, edge by edge, with colour-coded highlights."),
    ("**Display:**", "Rendering the optimal path as a glowing green 3D tube with performance statistics."),
]
for bold_part, rest in lifecycle:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(bold_part + " ")
    r.bold = True; r.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)

add_heading(doc, "Advanced Logic & Rendering", 2)

pillars = [
    ("**Adjacency List (JavaScript Map):**", "The graph is stored as Map<nodeId, [{to, weight}]>, enabling O(1) neighbour lookups. Each road's weight = road length + random traffic factor."),
    ("**Priority Queue Simulation:**", "Dijkstra's and A* use an array sorted by cost / f-score to always process the cheapest candidate next — simulating a min-heap."),
    ("**Admissible Heuristic (A*):**", "Euclidean 3D distance between two city positions serves as the A* heuristic. Since road distance ≥ straight-line distance, it never overestimates, guaranteeing optimal paths."),
    ("**Computation–Rendering Decoupling:**", "All algorithms run instantly in JavaScript. The exploration animation is a visual replay of the recorded edge trace using setInterval — the UI never freezes."),
    ("**3D Rendering via Three.js:**", "The scene uses WebGLRenderer, PerspectiveCamera, and OrbitControls. Roads are TubeGeometry wrapped around CatmullRomCurve3 splines. City labels are SpriteMaterial billboards that scale dynamically with zoom."),
]
for bold_part, rest in pillars:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(bold_part + " ")
    r.bold = True; r.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 3. REQUIREMENTS
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "3. Requirements", 1)

add_heading(doc, "Software Requirements", 2)

sw = [
    ("**Core Language:**", "JavaScript (ES Modules / ES2022) — engine for all graph logic, algorithms, and 3D interaction."),
    ("**3D Library:**", "Three.js (r160+) — WebGL-based 3D scene rendering, geometry, materials, and controls."),
    ("**Development Server:**", "npx serve via Node.js — required because ES Modules are blocked by browsers under the file:// protocol; runs at http://localhost:5500."),
    ("**Package Manager:**", "npm (Node.js) — used to launch the dev server via npm run dev."),
    ("**Browser:**", "Any modern browser with WebGL support (Chrome, Firefox, Edge, Safari)."),
]
for bold_part, rest in sw:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(bold_part + " ")
    r.bold = True; r.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)

add_heading(doc, "Key Modules", 2)

modules = [
    ("THREE.Scene",                       "The 3D world container"),
    ("THREE.PerspectiveCamera",           "The viewpoint and field of view"),
    ("THREE.WebGLRenderer",               "GPU-accelerated canvas rendering"),
    ("OrbitControls",                     "Mouse drag / zoom / pan interaction"),
    ("THREE.TubeGeometry + CatmullRomCurve3", "Curved 3D road meshes"),
    ("THREE.CanvasTexture",               "City name and weight label sprites"),
]

tbl = doc.add_table(rows=len(modules)+1, cols=2)
tbl.style = 'Table Grid'
headers = ["Module", "Purpose"]
for i, h in enumerate(headers):
    c = tbl.rows[0].cells[i]
    set_cell_bg(c, "1A568E")
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, (mod, purp) in enumerate(modules):
    row = tbl.rows[ri + 1]
    fill = "F5F9FC" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate([mod, purp]):
        c = row.cells[ci]
        set_cell_bg(c, fill)
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(10)
        if ci == 0:
            r.font.name = 'Courier New'

doc.add_paragraph()
add_heading(doc, "Hardware Requirements", 2)

hw = [
    ("**Processor:**", "Any modern dual-core CPU capable of running a WebGL context."),
    ("**RAM:**", "Minimum 2 GB system RAM; the browser tab itself uses under 200 MB."),
    ("**GPU:**", "Any integrated graphics chip with WebGL 1.0 support (released after 2010)."),
    ("**Storage:**", "Under 1 MB total — only 3 source files: index.html, main.js, package.json."),
    ("**Display:**", "1280×720 resolution or higher recommended for the dual-panel layout."),
]
for bold_part, rest in hw:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(bold_part + " ")
    r.bold = True; r.font.size = Pt(11)
    r2 = p.add_run(rest)
    r2.font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 4. CODE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "4. Code", 1)
add_heading(doc, "Core Concepts & DSA Summary: main.js", 2)

add_heading(doc, "1. Data Structures (DSA)", 3)

ds_items = [
    ("**Adjacency List (JavaScript Map — Hash Map):**",
     "The graph is stored as adjList, a Map where each key is a node ID and value is an array of {to, weight} neighbour objects. Provides O(1) neighbour lookup during traversal.",
     "adjList.get(cityA) → [\n  { to: cityB, weight: 22 },\n  { to: cityC, weight: 15 }\n]"),
    ("**Priority Queue (Array + Sort):**",
     "Dijkstra's and A* use a sorted array as a min-priority queue. Sorted by weight (Dijkstra) or f = g + h (A*) before each pop.",
     None),
    ("**Set (Hash Set):**",
     "Used in BFS and DFS as a visited set for O(1) membership checks — preventing nodes from being processed more than once.",
     None),
    ("**Stack (LIFO Array):**",
     "DFS uses .push() and .pop() to simulate a stack — diving deep into one branch before backtracking.",
     None),
    ("**Queue (FIFO Array):**",
     "BFS uses .push() and .shift() for level-by-level traversal.",
     None),
    ("**cameFrom Map:**",
     "All four algorithms use a cameFrom Map (nodeId → parentId) to record the shortest path tree, enabling O(n) path reconstruction.",
     None),
]
for bold_part, rest, code in ds_items:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(bold_part + " ")
    r.bold = True; r.font.size = Pt(11)
    p.add_run(rest).font.size = Pt(11)
    if code:
        add_code(doc, code)

add_heading(doc, "2. Algorithmic Techniques", 3)

add_para(doc, "**Dijkstra's Algorithm:** Explores the graph by always expanding the node with the lowest cumulative cost from the source. Guarantees the minimum-weight path.")
add_code(doc, """while (pq.length > 0) {
    pq.sort((a, b) => a.weight - b.weight);   // cheapest first
    const current = pq.shift().node;
    if (current === endNode) return reconstructPath(cameFrom, current);
    for (const neighbor of adjList.get(current)) {
        const alt = dist.get(current) + neighbor.weight;
        if (alt < dist.get(neighbor.to)) {
            dist.set(neighbor.to, alt);
            cameFrom.set(neighbor.to, current);
            pq.push({ node: neighbor.to, weight: alt });
        }
    }
}""")

add_para(doc, "**A* Search (Heuristic-guided):** Extends Dijkstra by adding a heuristic h(n) — the Euclidean 3D distance from node n to the destination. Priority queue is sorted by f = g + h.")
add_code(doc, """const h = (nodeId) => nodes[nodeId].pos.distanceTo(nodes[endNode].pos);
// f = g + h: actual cost + estimated remaining distance
fScore.set(neighbor.to, tentative_g + h(neighbor.to));
pq.sort((a, b) => a.f - b.f);""")

add_para(doc, "**BFS:** Traverses level by level using a FIFO queue, ignoring edge weights. Finds path with fewest hops.")
add_para(doc, "**DFS:** Traverses using a LIFO stack, diving deep before backtracking. Non-optimal — used for demonstration.")

add_para(doc, "**Path Reconstruction (reconstructPath):** Traces the cameFrom map backwards from destination to source.")
add_code(doc, """function reconstructPath(cameFrom, current) {
    const path = [current];
    while (cameFrom.has(current)) {
        current = cameFrom.get(current);
        path.unshift(current);
    }
    return path;
}""")

add_heading(doc, "3. Algorithm Comparison", 3)

comp_headers = ["Feature", "Dijkstra", "A*", "BFS", "DFS"]
comp_rows = [
    ["Optimal Path?",       "✅ (weight)", "✅ (weight)", "✅ (hops)", "❌ No"],
    ["Considers Weights?",  "✅",          "✅",           "❌",        "❌"],
    ["Uses Heuristic?",     "❌",          "✅",           "❌",        "❌"],
    ["Speed",               "Medium",      "Fast",         "Medium",    "Unpredictable"],
    ["Data Structure",      "Priority Queue","Priority Queue","Queue (FIFO)","Stack (LIFO)"],
    ["Best Use Case",       "Weighted shortest","Weighted + faster","Min hops","Exploration demo"],
]

tbl = doc.add_table(rows=len(comp_rows)+1, cols=5)
tbl.style = 'Table Grid'

for i, h in enumerate(comp_headers):
    c = tbl.rows[0].cells[i]
    set_cell_bg(c, "1A568E")
    r = c.paragraphs[0].add_run(h)
    r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

for ri, row_data in enumerate(comp_rows):
    row = tbl.rows[ri + 1]
    fill = "F5F9FC" if ri % 2 == 0 else "FFFFFF"
    for ci, val in enumerate(row_data):
        c = row.cells[ci]
        set_cell_bg(c, fill)
        r = c.paragraphs[0].add_run(val)
        r.font.size = Pt(10)
        if ci == 0:
            r.bold = True

add_heading(doc, "4. 3D Visualization & Rendering", 3)

vis = [
    ("**Terrain Generation:**", "A PlaneGeometry has its vertices displaced using sine/cosine math to create organic hills."),
    ("**City Nodes:**", "Each of 80 cities is a CylinderGeometry mesh. City names float as SpriteMaterial billboards, scaling dynamically with camera zoom."),
    ("**Road Edges:**", "Roads are TubeGeometry wrapped around a CatmullRomCurve3 spline, making them appear curved and organic."),
    ("**Search Animation:**", "After algorithm finishes, visitedEdges are replayed via setInterval (25ms delay), colouring each edge yellow."),
    ("**Final Path:**", "Drawn as a thicker tube (tubeRadius: 0.9) with emissive green material (emissive: 0x22c55e, emissiveIntensity: 1.5)."),
]
for bold_part, rest in vis:
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(bold_part + " ")
    r.bold = True; r.font.size = Pt(11)
    p.add_run(rest).font.size = Pt(11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
# 5. REAL-TIME EXAMPLE
# ══════════════════════════════════════════════════════════════════════════════
add_heading(doc, "5. Real-Time Example", 1)
add_para(doc, "This section illustrates how the system functions in a practical navigational scenario — a **logistics dispatcher** routing delivery vehicles across a city network.")

phases = [
    ("Phase 1: Route Setup (Graph Construction)",
     [("When the dispatcher opens the application in the browser (http://localhost:5500), the 3D terrain renders with all 80 cities (labeled cylinders) and their interconnecting road tubes."),
      ("The adjacency list is fully initialized. Every city knows its neighbours and the cost of reaching them."),
      ("Each road's weight reflects distance plus a randomized traffic factor — simulating real-world dynamic road conditions.")]),
    ("Phase 2: Finding the Optimal Route (The 'Dijkstra' Operation)",
     [("The dispatcher types 'Whiterun' in the Start bar and 'Solitude' in the Destination bar using the autocomplete search."),
      ("With Dijkstra selected, the system runs the full traversal instantly, recording every edge examined."),
      ("Yellow edges ripple outward from Whiterun across the map, showing every road the algorithm considered."),
      ("A glowing green tube traces the optimal route. The result modal shows: Total Distance, Nodes Visited, and Time Taken in ms.")]),
    ("Phase 3: Speed Comparison (The 'A*' Operation)",
     [("The dispatcher resets the map, selects the same start and destination, and switches to A* in the dropdown."),
      ("A* uses the straight-line 3D distance to Solitude as a guide, skipping cities clearly moving away from the destination."),
      ("The same optimal path is found, but the animation shows visibly fewer yellow edges — A* explored only ~24 nodes vs Dijkstra's 42."),
      ("Insight: Both found the same route. A* was smarter about where it looked.")]),
    ("Phase 4: Exploring Alternatives (BFS & DFS Demo)",
     [("BFS: The animation expands level by level, like a flood. The path has fewer city-hops but higher total weight — proving BFS ignores road costs."),
      ("DFS: The animation dives deep into one branch before backtracking erratically. The path is neither the shortest in hops nor in weight."),
      ("History Panel: All four searches are logged side by side with their costs and node counts for direct comparison.")]),
    ("Phase 5: Interactive Exploration (Camera & Labels)",
     [("Zoom In: The camera tilts from top-down toward a street-level view (GTA V-style dynamic tilt). City names scale down to remain readable."),
      ("Zoom to Cursor: Scrolling zooms toward wherever the mouse is pointing — the map slides toward the cursor, not the center."),
      ("Panel Toggle: The left control panel is collapsed via the '›' button, giving a full-screen view without losing the green path."),
      ("Reset: Clicking Reset Map clears all highlights, returning the scene to its original idle state.")]),
]

for phase_title, items in phases:
    add_heading(doc, phase_title, 2)
    for item in items:
        add_bullet(doc, item)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Project developed as part of the Data Structures and Algorithm course submission.")
r.italic = True
r.font.size = Pt(10)
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

# ── Save ──────────────────────────────────────────────────────────────────────
doc.save(r"d:\DSAPRO\threejs-dsa\InterCityPathfinder_Documentation.docx")
print("✅ InterCityPathfinder_Documentation.docx saved successfully!")
